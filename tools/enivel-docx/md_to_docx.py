#!/usr/bin/env python3
"""
Convert eNivel markdown specifications to DOCX with PDF-like styling.

- Real Heading 1/2/3/4 styles (blue #2c5fa1 for h1/h2)
- Real bullet/numbered lists
- Real tables with header row
- Real page headers (chapter title left, "eNivel" right) and footers (page N / total)
- Code blocks (monospace, gray shading)
- Mermaid blocks embedded as PNG images
- Inline `code` (monospace, orange background), **bold**, *italic*, [text](url)
- Blockquotes (indented, italic, blue accent)
- Manual page breaks at <!-- PAGEBREAK --> markers
"""

import json
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Pt, RGBColor, Emu
from docx.shared import Inches


BLUE = RGBColor(0x2C, 0x5F, 0xA1)
DARK = RGBColor(0x1F, 0x1F, 0x1F)
GRAY = RGBColor(0x6E, 0x6E, 0x6E)
ORANGE = RGBColor(0xB8, 0x42, 0x1A)


MERMAID_DIR = Path("/tmp/enivel-docx/mermaid")
with open(MERMAID_DIR / "manifest.json") as f:
    MERMAID_MANIFEST = json.load(f)


def set_cell_shading(cell, fill_hex):
    """Apply solid background color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, fill_hex):
    """Apply solid background color to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    p_pr.append(shd)


def add_page_number_field(paragraph):
    """Insert {PAGE} / {NUMPAGES} field into a paragraph for footer."""
    run1 = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText1 = OxmlElement("w:instrText")
    instrText1.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run1._r.append(fldChar1)
    run1._r.append(instrText1)
    run1._r.append(fldChar2)
    paragraph.add_run(" / ")
    run2 = paragraph.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run2._r.append(fldChar3)
    run2._r.append(instrText2)
    run2._r.append(fldChar4)


def add_hyperlink(paragraph, url, text):
    """Insert a hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2C5FA1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# Inline tokenizer for `code`, **bold**, *italic*, [link](url)
INLINE_RE = re.compile(
    r"(`[^`]+`)"               # code span
    r"|(\*\*[^*]+\*\*)"        # bold
    r"|(\*[^*]+\*)"            # italic
    r"|(\[[^\]]+\]\([^)]+\))"  # link
)


def add_inline_runs(paragraph, text, base_color=None):
    """Walk text, emitting runs for inline formatting."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos : m.start()])
            if base_color is not None:
                r.font.color.rgb = base_color
        token = m.group(0)
        if token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "SF Mono"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "SF Mono")
            r._element.rPr.rFonts.set(qn("w:hAnsi"), "SF Mono")
            r._element.rPr.rFonts.set(qn("w:cs"), "SF Mono")
            r.font.size = Pt(9)
            r.font.color.rgb = ORANGE
            # Shading (light cream)
            rPr = r._element.rPr
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F1EBE3")
            rPr.append(shd)
        elif token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
            if base_color is not None:
                r.font.color.rgb = base_color
        elif token.startswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
            if base_color is not None:
                r.font.color.rgb = base_color
        elif token.startswith("["):
            link_m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_m:
                label, url = link_m.group(1), link_m.group(2)
                add_hyperlink(paragraph, url, label)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        if base_color is not None:
            r.font.color.rgb = base_color


def configure_styles(doc):
    """Set up heading styles with PDF colors."""
    styles = doc.styles
    # Heading 1
    h1 = styles["Heading 1"]
    h1.font.name = "Helvetica"
    h1.font.color.rgb = BLUE
    h1.font.size = Pt(24)
    h1.font.bold = True
    # Heading 2
    h2 = styles["Heading 2"]
    h2.font.name = "Helvetica"
    h2.font.color.rgb = BLUE
    h2.font.size = Pt(16)
    h2.font.bold = True
    # Heading 3
    h3 = styles["Heading 3"]
    h3.font.name = "Helvetica"
    h3.font.color.rgb = DARK
    h3.font.size = Pt(13)
    h3.font.bold = True
    # Heading 4
    h4 = styles["Heading 4"]
    h4.font.name = "Helvetica"
    h4.font.color.rgb = DARK
    h4.font.size = Pt(11)
    h4.font.bold = True
    # Normal
    normal = styles["Normal"]
    normal.font.name = "Helvetica"
    normal.font.size = Pt(10)


def add_doc_meta(doc, date_str="2026-05-28"):
    p = doc.add_paragraph()
    r = p.add_run(f"eNivel-määrittely · Reaktor · {date_str}")
    r.font.color.rgb = GRAY
    r.font.size = Pt(9)


def setup_section_header_footer(section, header_left):
    """Configure section's header and footer."""
    # Header
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    # Use tabs to right-align "eNivel"
    tab_stops = header_p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17), alignment=2)  # right
    r1 = header_p.add_run(header_left)
    r1.font.size = Pt(8)
    r1.font.color.rgb = GRAY
    r1.font.name = "Helvetica"
    header_p.add_run("\t")
    r2 = header_p.add_run("eNivel")
    r2.font.size = Pt(8)
    r2.font.color.rgb = GRAY
    r2.font.name = "Helvetica"
    # Footer
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_p)
    for r in footer_p.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        r.font.name = "Helvetica"


def parse_table_rows(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i])
        i += 1
    # Parse cells
    parsed = []
    for r in rows:
        # Strip leading/trailing |
        r = r.strip()
        if r.startswith("|"):
            r = r[1:]
        if r.endswith("|"):
            r = r[:-1]
        cells = [c.strip() for c in r.split("|")]
        parsed.append(cells)
    # Filter alignment row (---)
    final = []
    for cells in parsed:
        if all(re.match(r"^:?-+:?$", c) for c in cells if c):
            continue
        final.append(cells)
    return final, i


def render_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row_cells in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell_p = cell.paragraphs[0]
            cell_p.text = ""  # clear default
            if c_idx < len(row_cells):
                add_inline_runs(cell_p, row_cells[c_idx])
            # Header row formatting
            if r_idx == 0:
                set_cell_shading(cell, "ECECEC")
                for r in cell_p.runs:
                    r.bold = True


def render_code_block(doc, code_lines):
    """Render a code block with monospace font and gray shading."""
    p = doc.add_paragraph()
    set_paragraph_shading(p, "F5F5F5")
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(code_lines):
        if i > 0:
            r = p.add_run()
            r.add_break(WD_BREAK.LINE)
        r = p.add_run(line)
        r.font.name = "SF Mono"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "SF Mono")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "SF Mono")
        r._element.rPr.rFonts.set(qn("w:cs"), "SF Mono")
        r.font.size = Pt(9)
        r.font.color.rgb = DARK


def render_blockquote(doc, text_lines):
    """Render a blockquote with left indent and italic."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_shading(p, "E8EEF5")
    text = " ".join(l.lstrip("> ").rstrip() for l in text_lines)
    add_inline_runs(p, text)
    for r in p.runs:
        r.italic = True


def render_mermaid(doc, slug, idx):
    """Insert a mermaid PNG image from manifest."""
    entries = MERMAID_MANIFEST.get(slug, [])
    if idx >= len(entries):
        p = doc.add_paragraph()
        r = p.add_run(f"[Mermaid-kaavio puuttuu: {slug} #{idx}]")
        r.italic = True
        r.font.color.rgb = GRAY
        return
    png_path = entries[idx]["path"]
    if not os.path.exists(png_path):
        p = doc.add_paragraph()
        r = p.add_run(f"[Mermaid-kuva puuttuu: {png_path}]")
        r.italic = True
        r.font.color.rgb = GRAY
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=Cm(15))


def convert_md_to_docx_body(doc, md_text, slug):
    """Walk markdown text, emitting DOCX content."""
    lines = md_text.split("\n")
    i = 0
    mermaid_counter = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()
        # Page break
        if "<!-- PAGEBREAK -->" in line:
            # Add a page break via paragraph
            p = doc.add_paragraph()
            r = p.add_run()
            r.add_break(WD_BREAK.PAGE)
            i += 1
            continue
        # Empty line
        if not stripped:
            i += 1
            continue
        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            # Strip MD link inline in headings (keep label)
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            p = doc.add_heading(level=min(level, 4))
            add_inline_runs(p, text)
            i += 1
            continue
        # Horizontal rule
        if re.match(r"^-{3,}\s*$", stripped):
            # In markdown specs these often separate index from body
            # Skip — page break above handles spacing
            i += 1
            continue
        # Code block fenced
        m = re.match(r"^```(\w*)\s*$", stripped)
        if m:
            lang = m.group(1)
            j = i + 1
            code_lines = []
            while j < len(lines) and not lines[j].rstrip() == "```":
                code_lines.append(lines[j])
                j += 1
            if lang == "mermaid":
                render_mermaid(doc, slug, mermaid_counter)
                mermaid_counter += 1
            else:
                render_code_block(doc, code_lines)
            i = j + 1
            continue
        # Blockquote
        if stripped.startswith(">"):
            j = i
            quote_lines = []
            while j < len(lines) and lines[j].strip().startswith(">"):
                quote_lines.append(lines[j])
                j += 1
            render_blockquote(doc, quote_lines)
            i = j
            continue
        # Table (starts with |)
        if stripped.startswith("|"):
            rows, j = parse_table_rows(lines, i)
            render_table(doc, rows)
            i = j
            continue
        # Bullet list (- or *)
        if re.match(r"^[-*]\s+", stripped):
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                content = re.sub(r"^[-*]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_inline_runs(p, content)
                i += 1
            continue
        # Numbered list (1. 2. ...)
        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                content = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_inline_runs(p, content)
                i += 1
            continue
        # Paragraph (possibly multi-line)
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            if not nxt:
                break
            # Stop if next line starts a new block
            if (
                re.match(r"^#{1,6}\s", nxt)
                or nxt.startswith("|")
                or re.match(r"^[-*]\s", nxt)
                or re.match(r"^\d+\.\s", nxt)
                or nxt.startswith(">")
                or re.match(r"^```", nxt)
                or "<!-- PAGEBREAK -->" in nxt
            ):
                break
            para_lines.append(nxt)
            j += 1
        text = " ".join(para_lines)
        p = doc.add_paragraph()
        add_inline_runs(p, text)
        i = j


def build_chapter_docx(md_path, header_left, out_path):
    doc = Document()
    # Page setup: A4 with margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    configure_styles(doc)
    setup_section_header_footer(section, header_left)
    add_doc_meta(doc)
    md = open(md_path, encoding="utf-8").read()
    slug = Path(md_path).stem
    convert_md_to_docx_body(doc, md, slug)
    doc.save(out_path)
    print(f"Wrote {out_path}")


def build_combined_docx(index_md, chapter_files, out_path):
    """Combined doc with index page + chapters as separate sections."""
    doc = Document()
    # First section: index page
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    configure_styles(doc)
    setup_section_header_footer(section, "eNivel-määrittely")
    add_doc_meta(doc)
    md = open(index_md, encoding="utf-8").read()
    slug = Path(index_md).stem
    convert_md_to_docx_body(doc, md, slug)
    # Append each chapter as new section
    for header_left, md_path in chapter_files:
        # Section break + new section
        new_section = doc.add_section()
        new_section.page_height = Cm(29.7)
        new_section.page_width = Cm(21.0)
        new_section.left_margin = Cm(1.8)
        new_section.right_margin = Cm(1.8)
        new_section.top_margin = Cm(2.0)
        new_section.bottom_margin = Cm(1.8)
        setup_section_header_footer(new_section, header_left)
        add_doc_meta(doc)
        md = open(md_path, encoding="utf-8").read()
        slug = Path(md_path).stem
        convert_md_to_docx_body(doc, md, slug)
    doc.save(out_path)
    print(f"Wrote {out_path}")


SPEC_DIR = Path("/Volumes/Evaka/enivel/specification")
LIITTEET_DIR = Path("/Volumes/Evaka/enivel/liiteet")
OUT_DIR = Path("/tmp/enivel-docx")
OUT_DIR.mkdir(parents=True, exist_ok=True)

CHAPTERS = [
    ("01-johdanto", "1. Johdanto", SPEC_DIR / "01-johdanto.md"),
    ("02-jarjestelmaarkitehtuuri", "2. Järjestelmäarkkitehtuuri", SPEC_DIR / "02-jarjestelmaarkitehtuuri.md"),
    ("03-tietomalli", "3. Tietomalli", SPEC_DIR / "03-tietomalli.md"),
    ("04-kayttotapaukset", "4. Käyttötapaukset ja vaatimukset", SPEC_DIR / "04-kayttotapaukset.md"),
    ("05-integraatiovaatimukset", "5. Integraatiovaatimukset", SPEC_DIR / "05-integraatiovaatimukset.md"),
    ("06-ei-funktionaaliset-vaatimukset", "6. Ei-funktionaaliset vaatimukset", SPEC_DIR / "06-ei-funktionaaliset-vaatimukset.md"),
    ("07-toteutus-nakokulmat", "7. Toteutuksen näkökulmat", SPEC_DIR / "07-toteutus-nakokulmat.md"),
]

LIITTEET = [
    ("A-tekniset-yksityiskohdat", "Liite A — Tekniset yksityiskohdat", LIITTEET_DIR / "A-tekniset-yksityiskohdat.md"),
    ("B-kayttoliittymaluonnokset", "Liite B — Käyttöliittymäluonnokset", LIITTEET_DIR / "B-kayttoliittymaluonnokset.md"),
]


def main():
    # Chapter-level docx
    for slug, header_left, md_path in CHAPTERS:
        build_chapter_docx(md_path, header_left, OUT_DIR / f"{slug}.docx")

    # Liite docx
    for slug, header_left, md_path in LIITTEET:
        build_chapter_docx(md_path, header_left, OUT_DIR / f"{slug}.docx")

    # Combined main spec
    build_combined_docx(
        SPEC_DIR / "enivel-specification.md",
        [(t, p) for _, t, p in CHAPTERS],
        OUT_DIR / "enivel-jarjestelmamaarittely-2026-05-28.docx",
    )

    # Combined liite
    build_combined_docx(
        LIITTEET_DIR / "enivel-liiteet.md",
        [(t, p) for _, t, p in LIITTEET],
        OUT_DIR / "enivel-liiteet-2026-05-28.docx",
    )


if __name__ == "__main__":
    main()
